import io
import zipfile
import pandas as pd
from PyPDF2 import PdfReader
from docx import Document


def extract_text_from_pdf(file_obj):
    pdf = PdfReader(file_obj)
    text = ""
    for page in pdf.pages:
        text += page.extract_text() or ""
    return text


def extract_text_from_docx(file_obj):
    doc = Document(file_obj)
    text = "\n".join([para.text for para in doc.paragraphs])
    return text


def extract_text_from_excel(file_obj):
    df = pd.read_excel(file_obj)
    text = ""
    for i, row in df.iterrows():
        row_text = ". ".join(f"{col}: {row[col]}" for col in df.columns)
        text += f"Row {i+1}: {row_text}\n"
    return text


def extract_text_from_zip(file_obj):
    text = ""
    with zipfile.ZipFile(file_obj) as z:
        for filename in z.namelist():
            with z.open(filename) as f:
                if filename.lower().endswith(".pdf"):
                    text += extract_text_from_pdf(f)
                elif filename.lower().endswith(".docx"):
                    text += extract_text_from_docx(f)
                elif filename.lower().endswith((".xlsx", ".xls")):
                    file_bytes = f.read()
                    text += extract_text_from_excel(io.BytesIO(file_bytes))
    return text


def get_raw_text(file_bytes: bytes, filename: str) -> str:
    raw_text = ""
    fname = filename.lower()

    if fname.endswith(".pdf"):
        pdf = PdfReader(io.BytesIO(file_bytes))
        for page in pdf.pages:
            raw_text += page.extract_text() or ""

    elif fname.endswith(".docx"):
        doc = Document(io.BytesIO(file_bytes))
        for para in doc.paragraphs:
            raw_text += para.text + "\n"

    elif fname.endswith((".xlsx", ".xls")):
        df = pd.read_excel(io.BytesIO(file_bytes))
        for i, row in df.iterrows():
            row_text = " | ".join(f"{col}: {row[col]}" for col in df.columns)
            raw_text += row_text + "\n"

    elif fname.endswith(".zip"):
        with zipfile.ZipFile(io.BytesIO(file_bytes)) as z:
            for inner_file in z.namelist():
                with z.open(inner_file) as f:
                    inner_bytes = f.read()
                    raw_text += get_raw_text(inner_bytes, inner_file)  # recursive

    return raw_text